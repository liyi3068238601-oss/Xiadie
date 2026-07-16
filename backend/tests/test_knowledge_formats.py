"""F.8 PDF/DOCX 注册解析、真实页码与格式错误测试。"""
import asyncio
import io
import json

import pytest

from app import db, knowledge, knowledge_chunker, knowledge_parser, knowledge_worker


@pytest.fixture(autouse=True)
def clean_knowledge_format_data():
    conn = db.connect()
    try:
        conn.execute("DELETE FROM knowledge_documents")
        conn.commit()
    finally:
        conn.close()
    for directory in (knowledge.STORAGE_DIR, knowledge.PARSED_DIR):
        directory.mkdir(parents=True, exist_ok=True)
        for path in directory.iterdir():
            if path.is_file():
                path.unlink()
    yield
    conn = db.connect()
    try:
        conn.execute("DELETE FROM knowledge_documents")
        conn.commit()
    finally:
        conn.close()
    for directory in (knowledge.STORAGE_DIR, knowledge.PARSED_DIR):
        for path in directory.iterdir():
            if path.is_file():
                path.unlink()


def _pdf(*pages: str) -> bytes:
    objects: list[bytes] = []
    page_ids = [3 + index * 2 for index in range(len(pages))]
    font_id = 3 + len(pages) * 2
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = " ".join(f"{value} 0 R" for value in page_ids).encode()
    objects.append(b"<< /Type /Pages /Kids [" + kids + b"] /Count " + str(len(pages)).encode() + b" >>")
    for index, text in enumerate(pages):
        page_id = page_ids[index]
        content_id = page_id + 1
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>".encode()
        )
        escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
        objects.append(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out.extend(f"{number} 0 obj\n".encode() + body + b"\nendobj\n")
    xref = len(out)
    out.extend(f"xref\n0 {len(objects)+1}\n".encode())
    out.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        out.extend(f"{offset:010d} 00000 n \n".encode())
    out.extend(f"trailer\n<< /Size {len(objects)+1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode())
    return bytes(out)


def _docx() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("角色资料", level=1)
    document.add_paragraph("遐蝶喜欢花朵与安静的星空。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "地点"
    table.cell(0, 1).text = "花园"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_pdf_parser_preserves_real_page_spans_and_chunks():
    raw = _pdf("First page has stars.", "Second page has flowers.")
    assert knowledge.validate_file("story.pdf", "application/pdf", raw)["decoded_chars"] == 0
    parsed = knowledge_parser.parse(raw, extension=".pdf")
    assert parsed["page_count"] == 2 and len(parsed["page_spans"]) == 2
    chunks = knowledge_chunker.chunk_artifact(parsed)
    assert chunks[0]["page_start"] == 1 and chunks[-1]["page_end"] == 2


def test_docx_parser_preserves_headings_paragraphs_and_tables_without_fake_pages():
    raw = _docx()
    knowledge.validate_file(
        "角色.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", raw,
    )
    parsed = knowledge_parser.parse(raw, extension=".docx")
    assert parsed["headings"][0]["title"] == "角色资料"
    assert "地点\t花园" in parsed["normalized_text"] and parsed["page_count"] == 0
    chunks = knowledge_chunker.chunk_artifact(parsed)
    assert json.loads(chunks[0]["heading_path_json"]) == ["角色资料"]
    assert all(chunk["page_start"] is None for chunk in chunks)


def test_corrupt_binary_formats_have_specific_errors():
    with pytest.raises(knowledge.KnowledgeImportError) as pdf:
        knowledge.validate_file("bad.pdf", "application/pdf", b"not a pdf")
    assert pdf.value.code == "pdf_signature_invalid"
    with pytest.raises(knowledge.KnowledgeImportError) as docx:
        knowledge.validate_file(
            "bad.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", b"PKbad",
        )
    assert docx.value.code == "docx_archive_invalid"
    with pytest.raises(knowledge_parser.ParserError) as empty_pdf:
        knowledge_parser.parse(_pdf(""), extension=".pdf")
    assert empty_pdf.value.code == "pdf_no_extractable_text"


def test_pdf_full_worker_pipeline_persists_page_count_and_locator():
    imported = knowledge.import_file("pages.pdf", "application/pdf", _pdf("Stars on page one.", "Flowers on page two."))
    assert asyncio.run(knowledge_worker.process_due(limit=3)) == 3
    document = knowledge.list_documents()[0]
    chunks = knowledge_worker.chunks_for_document(imported["document"]["id"])
    assert document["status"] == "indexed" and document["page_count"] == 2
    assert chunks[0]["page_start"] == 1 and chunks[-1]["page_end"] == 2
