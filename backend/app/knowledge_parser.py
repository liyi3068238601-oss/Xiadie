"""F.8 可注册的本地解析器：统一产出可验证文本、标题与真实页码跨度。"""
from __future__ import annotations

import hashlib
import io
import json
import re
from collections.abc import Callable

PARSER_VERSION = "knowledge-parser-registry-v2"
TEXT_PARSER_VERSION = "knowledge-text-parser-v2"
PDF_PARSER_VERSION = "knowledge-pdf-parser-v1"
DOCX_PARSER_VERSION = "knowledge-docx-parser-v1"
_HEADING = re.compile(r"^(#{1,6})[ \t]+(.+?)\s*$")
_FENCE = re.compile(r"^[ \t]{0,3}(`{3,}|~{3,})")


class ParserError(ValueError):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code


def parse(data: bytes, *, extension: str) -> dict:
    parser = _PARSERS.get(extension.casefold())
    if not parser:
        raise ParserError("parser_unsupported", "没有适用于该格式的本地解析器")
    return parser(data, extension.casefold())


def parser_version_for(extension: str) -> str:
    return {
        ".txt": TEXT_PARSER_VERSION, ".md": TEXT_PARSER_VERSION,
        ".pdf": PDF_PARSER_VERSION, ".docx": DOCX_PARSER_VERSION,
    }.get(extension.casefold(), "")


def artifact_bytes(result: dict) -> bytes:
    payload = {
        "parser_version": result["parser_version"],
        "normalized_text": result["normalized_text"],
        "normalized_sha256": result["normalized_sha256"],
        "char_count": result["char_count"], "line_count": result["line_count"],
        "headings": result["headings"], "page_count": result["page_count"],
        "page_spans": result["page_spans"],
    }
    return json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")


def _text_parser(data: bytes, extension: str) -> dict:
    try:
        text = data.decode("utf-8-sig", errors="strict")
    except UnicodeDecodeError as error:
        raise ParserError("encoding_unsupported", "文件不是有效的 UTF-8 文本") from error
    normalized = _normalize(text)
    headings = _markdown_headings(normalized) if extension == ".md" else []
    return _result(normalized, TEXT_PARSER_VERSION, headings, [])


def _pdf_parser(data: bytes, _extension: str) -> dict:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data), strict=True)
        if reader.is_encrypted:
            try:
                unlocked = reader.decrypt("")
            except Exception as error:  # noqa: BLE001
                raise ParserError("pdf_encrypted", "PDF 已加密，无法在本地读取") from error
            if not unlocked:
                raise ParserError("pdf_encrypted", "PDF 已加密，无法在本地读取")
        pages: list[str] = []
        for page in reader.pages:
            pages.append(_normalize(page.extract_text() or "").strip())
    except ParserError:
        raise
    except Exception as error:  # noqa: BLE001 - 对外只暴露稳定格式错误
        raise ParserError("pdf_parse_failed", "PDF 结构损坏或无法提取文本") from error
    if not any(pages):
        raise ParserError("pdf_no_extractable_text", "PDF 没有可提取文字，当前不执行 OCR")
    text, spans = _join_pages(pages)
    return _result(text, PDF_PARSER_VERSION, [], spans, page_count=len(pages))


def _docx_parser(data: bytes, _extension: str) -> dict:
    try:
        from docx import Document
        document = Document(io.BytesIO(data))
        blocks: list[tuple[str, int | None]] = []
        for paragraph in document.paragraphs:
            value = _normalize(paragraph.text).strip()
            if not value:
                continue
            style = (paragraph.style.name if paragraph.style else "") or ""
            match = re.match(r"(?i)^heading\s+([1-6])$", style.strip())
            blocks.append((value, int(match.group(1)) if match else None))
        for table in document.tables:
            for row in table.rows:
                value = "\t".join(_normalize(cell.text).strip() for cell in row.cells).strip()
                if value:
                    blocks.append((value, None))
    except Exception as error:  # noqa: BLE001
        raise ParserError("docx_parse_failed", "DOCX 结构损坏或无法提取文字") from error
    if not blocks:
        raise ParserError("docx_no_extractable_text", "DOCX 没有可提取文字")
    parts: list[str] = []
    headings: list[dict] = []
    line = 1
    for value, level in blocks:
        if level:
            headings.append({"level": level, "title": value, "line": line})
        parts.append(value)
        line += value.count("\n") + 2
    text = "\n\n".join(parts)
    return _result(text, DOCX_PARSER_VERSION, headings, [])


def _result(text: str, version: str, headings: list[dict], page_spans: list[dict],
            *, page_count: int = 0) -> dict:
    return {
        "parser_version": version, "normalized_text": text,
        "normalized_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
        "char_count": len(text), "line_count": 0 if not text else text.count("\n") + 1,
        "headings": headings, "heading_count": len(headings),
        "page_count": page_count, "page_spans": page_spans,
    }


def _join_pages(pages: list[str]) -> tuple[str, list[dict]]:
    parts: list[str] = []
    spans: list[dict] = []
    offset = 0
    for number, page in enumerate(pages, start=1):
        if parts:
            parts.append("\n\n")
            offset += 2
        start = offset
        parts.append(page)
        offset += len(page)
        spans.append({"page": number, "start": start, "end": offset})
    return "".join(parts), spans


def _normalize(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")


def _markdown_headings(text: str) -> list[dict]:
    headings: list[dict] = []
    fence: str | None = None
    for line_number, line in enumerate(text.splitlines(), start=1):
        fence_match = _FENCE.match(line)
        if fence_match:
            marker = fence_match.group(1)[0]
            if fence is None:
                fence = marker
            elif fence == marker:
                fence = None
            continue
        if fence is not None:
            continue
        match = _HEADING.match(line)
        if not match:
            continue
        title = re.sub(r"[ \t]+#+[ \t]*$", "", match.group(2)).strip()
        if title:
            headings.append({"level": len(match.group(1)), "title": title, "line": line_number})
    return headings


_PARSERS: dict[str, Callable[[bytes, str], dict]] = {
    ".txt": _text_parser, ".md": _text_parser,
    ".pdf": _pdf_parser, ".docx": _docx_parser,
}
